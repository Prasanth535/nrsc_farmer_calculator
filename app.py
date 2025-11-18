from flask import (
    Flask,
    render_template,
    request,
    redirect,
    url_for,
    session,
    send_file,
    jsonify,
)
from datetime import datetime, timedelta
import io
import os

from docx import Document

# ReportLab for PDF with border + tables
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors

import pandas as pd

# ================== NEW: DATABASE (NEON POSTGRESQL) SETUP ==================
from sqlalchemy import create_engine, text

# In Render, set DATABASE_URL env var to the Neon SQLAlchemy connection string
# Example:
# postgresql+psycopg2://neondb_owner:password@ep-xxxx.us-east-1.aws.neon.tech/neondb?sslmode=require
DATABASE_URL = os.environ.get("DATABASE_URL")

engine = None
if DATABASE_URL:
    try:
        # pool_pre_ping=True helps avoid stale connections
        engine = create_engine(DATABASE_URL, pool_pre_ping=True)
        # Optional: test connection once at startup
        with engine.connect() as conn:
            conn.execute(text("SELECT 1"))
        print("✅ Connected to Neon PostgreSQL successfully.")
    except Exception as e:
        print("⚠️ Error connecting to Neon PostgreSQL:", e)
else:
    print("⚠️ DATABASE_URL not set. Database connection will not be used.")

# ================== FLASK APP ==================
app = Flask(__name__)
app.secret_key = os.environ.get("SECRET_KEY", "CHANGE_THIS_SECRET_KEY")

# ================== FILE PATHS ==================
CROP_EXCEL_PATH = "crop_variety_stcr.xlsx"
STCR_CSV_PATH = "stcr_formulas.csv"

# ================== LOAD CROP–VARIETY DATA ==================
try:
    crop_df = pd.read_excel(CROP_EXCEL_PATH)
except Exception as e:
    print("Error loading Excel file:", e)
    crop_df = pd.DataFrame(
        columns=["Crop", "Variety", "Season", "Duration_days", "Rec_N", "Rec_P2O5", "Rec_K2O"]
    )

for col in ["Crop", "Variety", "Season"]:
    if col in crop_df.columns:
        crop_df[col] = crop_df[col].astype(str).str.strip()

if "Crop" in crop_df.columns:
    CROPS = sorted(crop_df["Crop"].dropna().unique().tolist())
else:
    CROPS = []

CROP_VARIETY_MAP = {}
if "Crop" in crop_df.columns and "Variety" in crop_df.columns:
    for crop_name in CROPS:
        vars_list = (
            crop_df.loc[crop_df["Crop"] == crop_name, "Variety"]
            .dropna()
            .unique()
            .tolist()
        )
        CROP_VARIETY_MAP[crop_name] = sorted(vars_list)


def get_crop_row(crop_name, variety_name):
    if crop_df.empty:
        return None
    crop_name = str(crop_name).strip()
    variety_name = str(variety_name).strip()
    sub = crop_df[
        (crop_df["Crop"] == crop_name) &
        (crop_df["Variety"] == variety_name)
    ]
    if sub.empty:
        sub = crop_df[crop_df["Crop"] == crop_name]
        if sub.empty:
            return None
    return sub.iloc[0].to_dict()

# ================== LOAD STCR FORMULAS ==================
try:
    stcr_df = pd.read_csv(STCR_CSV_PATH)
except Exception as e:
    print("Error loading STCR CSV file:", e)
    stcr_df = pd.DataFrame(
        columns=["Crop", "Variety", "TargetYield_t_ha",
                 "Formula_N", "Formula_P2O5", "Formula_K2O"]
    )

for col in ["Crop", "Variety"]:
    if col in stcr_df.columns:
        stcr_df[col] = stcr_df[col].astype(str).str.strip()


def get_stcr_row(crop_name, variety_name):
    if stcr_df.empty:
        return None

    crop_name = str(crop_name).strip()
    variety_name = str(variety_name).strip()

    exact = stcr_df[
        (stcr_df["Crop"] == crop_name) &
        (stcr_df["Variety"] == variety_name)
    ]
    if not exact.empty:
        return exact.iloc[0].to_dict()

    fallback = stcr_df[
        (stcr_df["Crop"] == crop_name) &
        (stcr_df["Variety"].str.upper().isin(["ALL", "", "NAN"]))
    ]
    if not fallback.empty:
        return fallback.iloc[0].to_dict()

    return None


def safe_eval_formula(formula_str, context):
    if not isinstance(formula_str, str) or not formula_str.strip():
        return 0.0

    allowed_names = {
        "T": context.get("T", 0.0),
        "SN": context.get("SN", 0.0),
        "SP": context.get("SP", 0.0),
        "SK": context.get("SK", 0.0),
        "max": max,
        "min": min,
    }
    try:
        value = eval(formula_str, {"__builtins__": {}}, allowed_names)
        return float(value)
    except Exception as e:
        print("Error evaluating formula:", formula_str, "Error:", e)
        return 0.0

# ================== TELANGANA SOIL TYPES ==================
TELANGANA_SOIL_TYPES = [
    "Red Clayey", "Red Loamy", "Red Shallow Loamy", "Red Shallow Clayey",
    "Red Shallow Gravelly Loam", "Red Shallow Gravelly Clayey",
    "Red Gravelly Loam", "Red Gravelly Clayey",
    "Red Calcareous Clayey", "Red Calcareous Gravelly Clayey",
    "Red Shallow Calcareous Gravelly Loam",
    "Medium Calcareous Black", "Deep Calcareous Black",
    "Deep Black", "Shallow Black", "Black (General)",
    "Red Shallow Gravelly Clay", "Red Shallow Gravelly Loamy",
    "Red Shallow Gravelly", "Red Gravelly",
    "Lateritic Gravelly Clayey", "Red Gravelly Clayey Loamy",
    "Alluvio Colluvial Clay", "Alluvio Colluvial Loamy",
    "Alluvio Colluvial Clayey", "Alluvio–Colluvial Clay",
    "Alluvio Colluvial Clay Loamy", "Alluvial Soil",
    "Alluvial Colluvial Loamy", "Alluvial Colluvial Clayey",
    "Saline Sodic Soil", "Calcareous Black",
    "Brown Forest Soils", "Rock Lands", "Water Body",
]

# ================== FERTILIZER SPLIT SCHEDULE ==================
FERT_SPLITS = {
    "Rice": [
        {"stage": "Basal",              "das_min": 0,  "das_max": 7,   "N_pct": 50, "P_pct": 100, "K_pct": 50},
        {"stage": "Tillering",          "das_min": 20, "das_max": 25,  "N_pct": 25, "P_pct": 0,   "K_pct": 25},
        {"stage": "Panicle initiation", "das_min": 40, "das_max": 45,  "N_pct": 25, "P_pct": 0,   "K_pct": 25},
    ],
    "Maize": [
        {"stage": "Basal",      "das_min": 0,  "das_max": 0,   "N_pct": 50, "P_pct": 100, "K_pct": 50},
        {"stage": "Knee-high",  "das_min": 25, "das_max": 30,  "N_pct": 25, "P_pct": 0,   "K_pct": 25},
        {"stage": "Tasseling",  "das_min": 45, "das_max": 50,  "N_pct": 25, "P_pct": 0,   "K_pct": 0},
    ],
    "Sorghum": [
        {"stage": "Basal",  "das_min": 0,  "das_max": 0,   "N_pct": 50, "P_pct": 100, "K_pct": 50},
        {"stage": "30 DAS", "das_min": 30, "das_max": 30,  "N_pct": 50, "P_pct": 0,   "K_pct": 50},
    ],
    "Wheat": [
        {"stage": "Basal",        "das_min": 0,  "das_max": 0,   "N_pct": 50, "P_pct": 100, "K_pct": 100},
        {"stage": "CRI",          "das_min": 20, "das_max": 22,  "N_pct": 25, "P_pct": 0,   "K_pct": 0},
        {"stage": "45–50 DAS",    "das_min": 45, "das_max": 50,  "N_pct": 25, "P_pct": 0,   "K_pct": 0},
    ],
    "Finger millet": [
        {"stage": "Basal",  "das_min": 0,  "das_max": 0,   "N_pct": 50, "P_pct": 100, "K_pct": 100},
        {"stage": "30 DAS", "das_min": 30, "das_max": 30,  "N_pct": 25, "P_pct": 0,   "K_pct": 0},
        {"stage": "45 DAS", "das_min": 45, "das_max": 45,  "N_pct": 25, "P_pct": 0,   "K_pct": 0},
    ],
    "Red gram": [
        {"stage": "Basal", "das_min": 0, "das_max": 0, "N_pct": 100, "P_pct": 100, "K_pct": 100},
    ],
    "Green gram": [
        {"stage": "Basal", "das_min": 0, "das_max": 0, "N_pct": 100, "P_pct": 100, "K_pct": 100},
    ],
    "Black gram": [
        {"stage": "Basal", "das_min": 0, "das_max": 0, "N_pct": 100, "P_pct": 100, "K_pct": 100},
    ],
    "Groundnut": [
        {"stage": "Basal", "das_min": 0, "das_max": 0, "N_pct": 100, "P_pct": 100, "K_pct": 100},
    ],
    "Soybean": [
        {"stage": "Basal", "das_min": 0, "das_max": 0, "N_pct": 100, "P_pct": 100, "K_pct": 100},
    ],
    "Cotton": [
        {"stage": "Basal",  "das_min": 0,  "das_max": 0,   "N_pct": 25, "P_pct": 100, "K_pct": 50},
        {"stage": "30 DAS", "das_min": 30, "das_max": 30,  "N_pct": 25, "P_pct": 0,   "K_pct": 25},
        {"stage": "60 DAS", "das_min": 60, "das_max": 60,  "N_pct": 25, "P_pct": 0,   "K_pct": 25},
        {"stage": "90 DAS", "das_min": 90, "das_max": 90,  "N_pct": 25, "P_pct": 0,   "K_pct": 0},
    ],
    "Chilli": [
        {"stage": "Basal",       "das_min": 0,  "das_max": 0,   "N_pct": 30, "P_pct": 100, "K_pct": 50},
        {"stage": "30 DAS",      "das_min": 30, "das_max": 30,  "N_pct": 25, "P_pct": 0,   "K_pct": 25},
        {"stage": "Flowering",   "das_min": 45, "das_max": 50,  "N_pct": 25, "P_pct": 0,   "K_pct": 25},
        {"stage": "Fruit set",   "das_min": 60, "das_max": 70,  "N_pct": 20, "P_pct": 0,   "K_pct": 0},
    ],
    "Turmeric": [
        {"stage": "Basal",  "das_min": 0,  "das_max": 0,   "N_pct": 50, "P_pct": 100, "K_pct": 100},
        {"stage": "45 DAS", "das_min": 45, "das_max": 45,  "N_pct": 25, "P_pct": 0,   "K_pct": 0},
        {"stage": "90 DAS", "das_min": 90, "das_max": 90,  "N_pct": 25, "P_pct": 0,   "K_pct": 0},
    ],
}

# ================== IRRIGATION DEPTH (mm) BY CROP ==================
CROP_IRRIGATION_DEPTH_MM = {
    "Rice": 50,
    "Maize": 40,
    "Wheat": 45,
    "Cotton": 35,
    "Chilli": 30,
    "Turmeric": 35,
    "Groundnut": 30,
}
DEFAULT_IRRIGATION_DEPTH_MM = 40

# ================== HELPERS ==================
def compute_averages(readings):
    if not readings:
        return {}
    keys = readings[0].keys()
    averages = {}
    for k in keys:
        try:
            averages[k] = sum(float(r.get(k, 0.0) or 0.0) for r in readings) / len(readings)
        except Exception:
            averages[k] = 0.0
    return averages


def generate_irrigation_schedule(sowing_date_str, duration_days, crop_name, area_ha):
    """Weekly irrigations with fixed depth per crop; returns list of dicts."""
    if not sowing_date_str or not duration_days:
        return []
    try:
        sowing_date = datetime.strptime(sowing_date_str, "%Y-%m-%d").date()
    except ValueError:
        return []
    try:
        duration_days = int(duration_days)
    except Exception:
        duration_days = 0

    depth_mm = CROP_IRRIGATION_DEPTH_MM.get(crop_name, DEFAULT_IRRIGATION_DEPTH_MM)
    # 1 mm of water on 1 ha = 10 m3
    water_per_ha_m3 = depth_mm * 10.0
    water_field_m3 = water_per_ha_m3 * area_ha

    schedule = []
    d = sowing_date + timedelta(days=7)
    i = 1
    while (d - sowing_date).days <= duration_days:
        schedule.append(
            {
                "no": i,
                "date": d.isoformat(),
                "depth_mm": float(depth_mm),
                "water_m3_per_ha": float(water_per_ha_m3),
                "water_m3_field": float(water_field_m3),
            }
        )
        d += timedelta(days=7)
        i += 1
    return schedule


def soil_type_factor(soil_type):
    if not soil_type:
        return 1.0
    st = soil_type.lower()
    if "deep black" in st or "calcareous" in st:
        return 0.9
    if "saline" in st or "sodic" in st:
        return 0.85
    if "shallow" in st or "gravelly" in st:
        return 1.0
    return 1.0


def combo_urea_dap_mop(n_kg, p2o5_kg, k2o_kg):
    """Return (urea, dap, mop) in kg/ha for given N, P2O5, K2O requirement."""
    n_kg = max(0.0, float(n_kg or 0.0))
    p2o5_kg = max(0.0, float(p2o5_kg or 0.0))
    k2o_kg = max(0.0, float(k2o_kg or 0.0))

    dap = p2o5_kg / 0.46 if p2o5_kg > 0 else 0.0
    n_from_dap = dap * 0.18
    n_remaining = max(0.0, n_kg - n_from_dap)
    urea = n_remaining / 0.46 if n_remaining > 0 else 0.0
    mop = k2o_kg / 0.60 if k2o_kg > 0 else 0.0
    return urea, dap, mop


def combo_urea_ssp_mop(n_kg, p2o5_kg, k2o_kg):
    """Return (urea, ssp, mop) in kg/ha for given N, P2O5, K2O requirement."""
    n_kg = max(0.0, float(n_kg or 0.0))
    p2o5_kg = max(0.0, float(p2o5_kg or 0.0))
    k2o_kg = max(0.0, float(k2o_kg or 0.0))

    ssp = p2o5_kg / 0.16 if p2o5_kg > 0 else 0.0
    urea = n_kg / 0.46 if n_kg > 0 else 0.0
    mop = k2o_kg / 0.60 if k2o_kg > 0 else 0.0
    return urea, ssp, mop


def npk_to_fertilizer_form(n_kg, p2o5_kg, k2o_kg):
    """Simple text summary (Urea + DAP + MOP per ha)."""
    urea, dap, mop = combo_urea_dap_mop(n_kg, p2o5_kg, k2o_kg)
    return (
        f"Urea: {urea:.1f} kg/ha, "
        f"DAP: {dap:.1f} kg/ha, "
        f"MOP: {mop:.1f} kg/ha"
    )


def generate_fertilizer_schedule(crop_name, sowing_date_str, n_total, p_total, k_total, area_ha=1.0):
    if not sowing_date_str:
        return []
    try:
        sowing_date = datetime.strptime(sowing_date_str, "%Y-%m-%d").date()
    except ValueError:
        return []
    splits = FERT_SPLITS.get(crop_name, [])
    if not splits:
        splits = [
            {"stage": "Basal", "das_min": 0, "das_max": 0,
             "N_pct": 100, "P_pct": 100, "K_pct": 100}
        ]
    schedule = []
    for s in splits:
        das_min = s.get("das_min", 0) or 0
        das_apply = 0 if das_min == 0 else das_min
        app_date = sowing_date + timedelta(days=das_apply)

        N_pct = s.get("N_pct", 0) or 0
        P_pct = s.get("P_pct", 0) or 0
        K_pct = s.get("K_pct", 0) or 0

        n_amt = round(n_total * N_pct / 100.0, 2) if N_pct > 0 else 0.0
        p_amt = round(p_total * P_pct / 100.0, 2) if P_pct > 0 else 0.0
        k_amt = round(k_total * K_pct / 100.0, 2) if K_pct > 0 else 0.0

        ud_urea, ud_dap, ud_mop = combo_urea_dap_mop(n_amt, p_amt, k_amt)
        us_urea, us_ssp, us_mop = combo_urea_ssp_mop(n_amt, p_amt, k_amt)

        schedule.append(
            {
                "stage": s.get("stage", ""),
                "das": das_apply,
                "date": app_date.isoformat(),
                "N": n_amt,
                "P2O5": p_amt,
                "K2O": k_amt,
                # Urea + DAP + MOP (per ha and per field)
                "urea_dap_urea_ha": round(ud_urea, 2),
                "urea_dap_dap_ha": round(ud_dap, 2),
                "urea_dap_mop_ha": round(ud_mop, 2),
                "urea_dap_urea_field": round(ud_urea * area_ha, 2),
                "urea_dap_dap_field": round(ud_dap * area_ha, 2),
                "urea_dap_mop_field": round(ud_mop * area_ha, 2),
                # Urea + SSP + MOP (per ha and per field)
                "urea_ssp_urea_ha": round(us_urea, 2),
                "urea_ssp_ssp_ha": round(us_ssp, 2),
                "urea_ssp_mop_ha": round(us_mop, 2),
                "urea_ssp_urea_field": round(us_urea * area_ha, 2),
                "urea_ssp_ssp_field": round(us_ssp * area_ha, 2),
                "urea_ssp_mop_field": round(us_mop * area_ha, 2),
            }
        )
    return schedule

# ================== ROUTES ==================
@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        lang = request.form.get("language", "en")
        session["language"] = lang
        return redirect(url_for("farmer_details"))
    return render_template("index.html")


@app.route("/farmer-details", methods=["GET", "POST"])
def farmer_details():
    if request.method == "POST"]:
        farmer = {
            "name": request.form.get("name", "").strip(),
            "mobile": request.form.get("mobile", "").strip(),
            "village": request.form.get("village", "").strip(),
            "mandal": request.form.get("mandal", "").strip(),
            "district": request.form.get("district", "").strip(),
        }
        session["farmer"] = farmer
        return redirect(url_for("crop_details"))

    farmer = session.get("farmer", {})
    return render_template("farmer_details.html", farmer=farmer)


@app.route("/crop-details", methods=["GET", "POST"])
def crop_details():
    crop_session = session.get("crop", {})

    if request.method == "POST":
        selected_crop = request.form.get("crop") or ""
        selected_variety = request.form.get("variety") or ""
        soil_type = request.form.get("soil_type") or ""
        gps_location = request.form.get("gps_location", "").strip()
        gps_address = request.form.get("gps_address", "").strip()
        sowing_date = request.form.get("sowing_date", "").strip()

        land_area_str = request.form.get("land_area", "").strip()
        try:
            land_area = float(land_area_str) if land_area_str else 1.0
        except ValueError:
            land_area = 1.0

        ty_str = request.form.get("target_yield_q", "").strip()
        try:
            target_yield_q = float(ty_str) if ty_str else None
        except ValueError:
            target_yield_q = None

        # convert q/ha to t/ha for STCR formulas
        target_yield_t = target_yield_q / 10.0 if target_yield_q is not None else None

        row = get_crop_row(selected_crop, selected_variety)

        season = ""
        duration = 0
        rec_n = rec_p = rec_k = 0.0

        if row:
            season = row.get("Season", "") or ""
            try:
                duration = int(row.get("Duration_days", 0) or 0)
            except Exception:
                duration = 0
            rec_n = float(row.get("Rec_N", 0) or 0)
            rec_p = float(row.get("Rec_P2O5", 0) or 0)
            rec_k = float(row.get("Rec_K2O", 0) or 0)
        else:
            duration_str = request.form.get("duration", "").strip()
            try:
                duration = int(duration_str) if duration_str else 0
            except ValueError:
                duration = 0

        recommended_npk_str = f"{rec_n}:{rec_p}:{rec_k}"

        crop_session = {
            "gps_location": gps_location,
            "gps_address": gps_address,
            "soil_type": soil_type,
            "crop": selected_crop,
            "variety": selected_variety,
            "season": season,
            "duration": duration,
            "recommended_npk": recommended_npk_str,
            "sowing_date": sowing_date,
            "land_area": land_area,
            "target_yield_q": target_yield_q,
            "target_yield_t": target_yield_t,
        }

        session["crop"] = crop_session
        session["stcr_rec"] = {"N": rec_n, "P2O5": rec_p, "K2O": rec_k}
        session["readings"] = []

        return redirect(url_for("readings"))

    return render_template(
        "crop_details.html",
        crop=crop_session,
        crops=CROPS,
        crop_variety_map=CROP_VARIETY_MAP,
        telangana_soils=TELANGANA_SOIL_TYPES,
    )


@app.route("/api/crop-info")
def crop_info_api():
    crop_name = request.args.get("crop", "").strip()
    variety_name = request.args.get("variety", "").strip()
    row = get_crop_row(crop_name, variety_name)
    if not row:
        return jsonify({"ok": False}), 404

    try:
        duration = int(row.get("Duration_days", 0) or 0)
    except Exception:
        duration = 0

    rec_n = float(row.get("Rec_N", 0) or 0)
    rec_p = float(row.get("Rec_P2O5", 0) or 0)
    rec_k = float(row.get("Rec_K2O", 0) or 0)

    return jsonify(
        {
            "ok": True,
            "season": row.get("Season", "") or "",
            "duration": duration,
            "rec_n": rec_n,
            "rec_p": rec_p,
            "rec_k": rec_k,
            "recommended_npk": f"{rec_n}:{rec_p}:{rec_k}",
        }
    )

# ---------- Readings page (GET only) ----------
@app.route("/readings", methods=["GET"])
def readings():
    readings_list = session.get("readings", [])
    reading_no = len(readings_list) + 1
    if reading_no > 10:
        reading_no = 10

    return render_template(
        "readings.html",
        readings=readings_list,
        reading_no=reading_no,
        max_readings=10,
    )

# ---------- AJAX API: save one reading ----------
@app.route("/api/save-reading", methods=["POST"])
def api_save_reading():
    data = request.get_json() or {}

    def get_float(name):
        val = data.get(name, 0)
        try:
            return float(val) if val not in [None, ""] else 0.0
        except (TypeError, ValueError):
            return 0.0

    reading = {
        "soil_n": get_float("soil_n"),
        "soil_p": get_float("soil_p"),
        "soil_k": get_float("soil_k"),
        "ec": get_float("ec"),
        "soil_temp": get_float("soil_temp"),
        "ph": get_float("ph"),
        "soil_moisture": get_float("soil_moisture"),
        "air_temp": get_float("air_temp"),
        "air_humidity": get_float("air_humidity"),
        "tds": get_float("tds"),
    }

    readings_list = session.get("readings", [])
    readings_list.append(reading)
    session["readings"] = readings_list

    done = len(readings_list) >= 10
    return jsonify({"ok": True, "count": len(readings_list), "done": done})

# ---------- AJAX API: reset readings ----------
@app.route("/api/reset-readings", methods=["POST"])
def api_reset_readings():
    session["readings"] = []
    return jsonify({"ok": True})

# ---------- Calculate & Report ----------
@app.route("/calculate")
def calculate():
    readings_list = session.get("readings", [])
    averages = compute_averages(readings_list)
    session["averages"] = averages

    crop = session.get("crop", {})
    stcr_rec = session.get("stcr_rec", {})
    soil_type = crop.get("soil_type")
    land_area = float(crop.get("land_area", 1.0) or 1.0)

    rec_n = float(stcr_rec.get("N", 0.0))
    rec_p = float(stcr_rec.get("P2O5", 0.0))
    rec_k = float(stcr_rec.get("K2O", 0.0))

    final_n = rec_n
    final_p = rec_p
    final_k = rec_k

    avg_nsoil = float(averages.get("soil_n", 0.0) or 0.0)
    avg_psoil = float(averages.get("soil_p", 0.0) or 0.0)
    avg_ksoil = float(averages.get("soil_k", 0.0) or 0.0)

    stcr_row = get_stcr_row(crop.get("crop"), crop.get("variety"))

    if stcr_row:
        target_from_farmer_t = crop.get("target_yield_t")
        if target_from_farmer_t is not None:
            T = float(target_from_farmer_t)
        else:
            try:
                T = float(stcr_row.get("TargetYield_t_ha", 0.0) or 0.0)
            except Exception:
                T = 0.0

        context = {"T": T, "SN": avg_nsoil, "SP": avg_psoil, "SK": avg_ksoil}

        stcr_n = safe_eval_formula(stcr_row.get("Formula_N", ""), context)
        stcr_p = safe_eval_formula(stcr_row.get("Formula_P2O5", ""), context)
        stcr_k = safe_eval_formula(stcr_row.get("Formula_K2O", ""), context)

        factor = soil_type_factor(soil_type)
        stcr_n *= factor
        stcr_p *= factor
        stcr_k *= factor

        stcr_n = max(0.0, stcr_n)
        stcr_p = max(0.0, stcr_p)
        stcr_k = max(0.0, stcr_k)

        final_n = min(rec_n, stcr_n) if rec_n > 0 else stcr_n
        final_p = min(rec_p, stcr_p) if rec_p > 0 else stcr_p
        final_k = min(rec_k, stcr_k) if rec_k > 0 else stcr_k

    recommended_npk_after = f"{final_n:.2f}:{final_p:.2f}:{final_k:.2f}"
    fertilizer_form = npk_to_fertilizer_form(final_n, final_p, final_k)

    # Fertilizer combinations per ha
    urea_dap_urea_ha, urea_dap_dap_ha, urea_dap_mop_ha = combo_urea_dap_mop(final_n, final_p, final_k)
    urea_ssp_urea_ha, urea_ssp_ssp_ha, urea_ssp_mop_ha = combo_urea_ssp_mop(final_n, final_p, final_k)

    calculated = {
        "recommended_npk_after": recommended_npk_after,
        "fertilizer_form": fertilizer_form,
        "final_n": final_n,
        "final_p": final_p,
        "final_k": final_k,
        "area_ha": land_area,
        "urea_dap_mop_per_ha": {
            "urea": round(urea_dap_urea_ha, 2),
            "dap": round(urea_dap_dap_ha, 2),
            "mop": round(urea_dap_mop_ha, 2),
        },
        "urea_dap_mop_field": {
            "urea": round(urea_dap_urea_ha * land_area, 2),
            "dap": round(urea_dap_dap_ha * land_area, 2),
            "mop": round(urea_dap_mop_ha * land_area, 2),
        },
        "urea_ssp_mop_per_ha": {
            "urea": round(urea_ssp_urea_ha, 2),
            "ssp": round(urea_ssp_ssp_ha, 2),
            "mop": round(urea_ssp_mop_ha, 2),
        },
        "urea_ssp_mop_field": {
            "urea": round(urea_ssp_urea_ha * land_area, 2),
            "ssp": round(urea_ssp_ssp_ha * land_area, 2),
            "mop": round(urea_ssp_mop_ha * land_area, 2),
        },
    }
    session["calculated"] = calculated

    irrigation_schedule = generate_irrigation_schedule(
        crop.get("sowing_date"),
        crop.get("duration", 0),
        crop.get("crop"),
        land_area,
    )
    session["irrigation_schedule"] = irrigation_schedule

    fert_schedule = generate_fertilizer_schedule(
        crop.get("crop"),
        crop.get("sowing_date"),
        final_n,
        final_p,
        final_k,
        land_area,
    )
    session["fert_schedule"] = fert_schedule

    return redirect(url_for("report"))


@app.route("/report")
def report():
    farmer = session.get("farmer", {})
    crop = session.get("crop", {})
    readings_list = session.get("readings", [])
    averages = session.get("averages", {})
    calculated = session.get("calculated", {})
    irrigation_schedule = session.get("irrigation_schedule", [])
    fert_schedule = session.get("fert_schedule", [])

    return render_template(
        "report.html",
        farmer=farmer,
        crop=crop,
        readings=readings_list,
        averages=averages,
        calculated=calculated,
        irrigation_schedule=irrigation_schedule,
        fert_schedule=fert_schedule,
    )

# ---------- Download as PDF ----------
@app.route("/download/pdf")
def download_pdf():
    farmer = session.get("farmer", {})
    crop = session.get("crop", {})
    averages = session.get("averages", {})
    calculated = session.get("calculated", {})
    irrigation_schedule = session.get("irrigation_schedule", [])
    fert_schedule = session.get("fert_schedule", [])

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=40,
        rightMargin=40,
        topMargin=60,
        bottomMargin=40,
    )
    styles = getSampleStyleSheet()
    elements = []

    # Title
    elements.append(Paragraph("NRSC Farmer Advisor - Soil & Crop Report", styles["Title"]))
    elements.append(Spacer(1, 12))

    # Farmer details table
    farmer_data = [
        ["Field", "Value"],
        ["Farmer Name", farmer.get("name", "")],
        ["Mobile", farmer.get("mobile", "")],
        ["Village", farmer.get("village", "")],
        ["Mandal", farmer.get("mandal", "")],
        ["District", farmer.get("district", "")],
    ]
    elements.append(Paragraph("Farmer Details", styles["Heading2"]))

    t = Table(farmer_data, colWidths=[120, 350])
    t.setStyle(TableStyle([
        ("GRID", (0,0), (-1,-1), 0.5, colors.grey),
        ("BACKGROUND", (0,0), (-1,0), colors.lightgrey),
        ("VALIGN", (0,0), (-1,-1), "MIDDLE"),
    ]))
    elements.append(t)
    elements.append(Spacer(1, 12))

    # Crop & soil table
    area_ha = float(crop.get("land_area", 1.0) or 1.0)
    ty_q = crop.get("target_yield_q")
    crop_data = [
        ["Field", "Value"],
        ["GPS Location", crop.get("gps_location", "")],
        ["Address", crop.get("gps_address", "")],
        ["Soil Type", crop.get("soil_type", "")],
        ["Crop", crop.get("crop", "")],
        ["Variety", crop.get("variety", "")],
        ["Season", crop.get("season", "")],
        ["Crop Duration (days)", str(crop.get("duration", ""))],
        ["Recommended N:P:K (kg/ha)", crop.get("recommended_npk", "")],
        ["Field / Land Area (ha)", f"{area_ha:.2f}"],
        ["Date of Sowing", crop.get("sowing_date", "")],
    ]
    if ty_q is not None:
        crop_data.insert(-1, ["Target yield for STCR (q/ha)", str(ty_q)])

    elements.append(Paragraph("Crop & Soil Details", styles["Heading2"]))
    t = Table(crop_data, colWidths=[170, 300])
    t.setStyle(TableStyle([
        ("GRID", (0,0), (-1,-1), 0.5, colors.grey),
        ("BACKGROUND", (0,0), (-1,0), colors.lightgrey),
        ("VALIGN", (0,0), (-1,-1), "MIDDLE"),
    ]))
    elements.append(t)
    elements.append(Spacer(1, 12))

    # Average sensor readings
    avg_data = [["Parameter", "Average value"]]
    for label, key in [
        ("Soil Nitrogen (N)", "soil_n"),
        ("Soil Phosphorus (P)", "soil_p"),
        ("Soil Potassium (K)", "soil_k"),
        ("Electrical Conductivity (EC)", "ec"),
        ("Soil Temperature (°C)", "soil_temp"),
        ("Soil pH", "ph"),
        ("Soil Moisture (%)", "soil_moisture"),
        ("Air Temperature (°C)", "air_temp"),
        ("Air Humidity (%)", "air_humidity"),
        ("Water/Soil TDS", "tds"),
    ]:
        avg_data.append([label, f"{averages.get(key, 0):.2f}"])
    elements.append(Paragraph("Average Sensor Readings", styles["Heading2"]))
    t = Table(avg_data, colWidths=[220, 250])
    t.setStyle(TableStyle([
        ("GRID", (0,0), (-1,-1), 0.5, colors.grey),
        ("BACKGROUND", (0,0), (-1,0), colors.lightgrey),
    ]))
    elements.append(t)
    elements.append(Spacer(1, 12))

    # Fertilizer recommendation summary
    elements.append(Paragraph("Fertilizer Recommendation (STCR-based)", styles["Heading2"]))
    elements.append(Paragraph(
        f"Final N:P:K recommendation (kg/ha): {calculated.get('recommended_npk_after', '')}",
        styles["Normal"],
    ))
    elements.append(Spacer(1, 6))

    # Tables for fertilizer combinations
    ud_ha = calculated.get("urea_dap_mop_per_ha", {})
    ud_field = calculated.get("urea_dap_mop_field", {})
    us_ha = calculated.get("urea_ssp_mop_per_ha", {})
    us_field = calculated.get("urea_ssp_mop_field", {})

    # Urea + DAP + MOP
    data_ud = [
        ["Basis", "Urea (kg)", "DAP (kg)", "MoP (kg)"],
        ["Per hectare",
         f"{ud_ha.get('urea', 0):.2f}",
         f"{ud_ha.get('dap', 0):.2f}",
         f"{ud_ha.get('mop', 0):.2f}"],
        [f"For field ({area_ha:.2f} ha)",
         f"{ud_field.get('urea', 0):.2f}",
         f"{ud_field.get('dap', 0):.2f}",
         f"{ud_field.get('mop', 0):.2f}"],
    ]
    elements.append(Paragraph("Option 1: Urea + DAP + MoP", styles["Heading3"]))
    t = Table(data_ud, colWidths=[140, 90, 90, 90])
    t.setStyle(TableStyle([
        ("GRID", (0,0), (-1,-1), 0.5, colors.grey),
        ("BACKGROUND", (0,0), (-1,0), colors.lightgrey),
    ]))
    elements.append(t)
    elements.append(Spacer(1, 8))

    # Urea + SSP + MOP
    data_us = [
        ["Basis", "Urea (kg)", "SSP (kg)", "MoP (kg)"],
        ["Per hectare",
         f"{us_ha.get('urea', 0):.2f}",
         f"{us_ha.get('ssp', 0):.2f}",
         f"{us_ha.get('mop', 0):.2f}"],
        [f"For field ({area_ha:.2f} ha)",
         f"{us_field.get('urea', 0):.2f}",
         f"{us_field.get('ssp', 0):.2f}",
         f"{us_field.get('mop', 0):.2f}"],
    ]
    elements.append(Paragraph("Option 2: Urea + SSP + MoP", styles["Heading3"]))
    t = Table(data_us, colWidths=[140, 90, 90, 90])
    t.setStyle(TableStyle([
        ("GRID", (0,0), (-1,-1), 0.5, colors.grey),
        ("BACKGROUND", (0,0), (-1,0), colors.lightgrey),
    ]))
    elements.append(t)
    elements.append(Spacer(1, 12))

    # Fertilizer split schedule
    if fert_schedule:
        elements.append(Paragraph("Fertilizer Split Schedule (N, P₂O₅, K₂O in kg/ha)", styles["Heading2"]))

        data_fs = [["Stage", "DAS", "Date", "N", "P₂O₅", "K₂O"]]
        for fs in fert_schedule:
            data_fs.append([
                fs["stage"],
                str(fs["das"]),
                fs["date"],
                f"{fs['N']:.2f}",
                f"{fs['P2O5']:.2f}",
                f"{fs['K2O']:.2f}",
            ])
        t = Table(data_fs, colWidths=[110, 40, 80, 60, 60, 60])
        t.setStyle(TableStyle([
            ("GRID", (0,0), (-1,-1), 0.5, colors.grey),
            ("BACKGROUND", (0,0), (-1,0), colors.lightgrey),
        ]))
        elements.append(t)
        elements.append(Spacer(1, 8))

        elements.append(Paragraph(
            "Equivalent fertilizer quantities (for your field) at each stage "
            "are provided in the HTML report for detailed reference.",
            styles["Italic"],
        ))
        elements.append(Spacer(1, 12))

    # Irrigation schedule table
    if irrigation_schedule:
        elements.append(Paragraph("Irrigation Schedule", styles["Heading2"]))
        data_ir = [["No.", "Date", "Depth (mm)", "Water (m³/ha)", "Water (m³) for field"]]
        for irr in irrigation_schedule:
            data_ir.append([
                str(irr["no"]),
                irr["date"],
                f"{irr['depth_mm']:.1f}",
                f"{irr['water_m3_per_ha']:.1f}",
                f"{irr['water_m3_field']:.1f}",
            ])
        t = Table(data_ir, colWidths=[35, 80, 70, 90, 110])
        t.setStyle(TableStyle([
            ("GRID", (0,0), (-1,-1), 0.5, colors.grey),
            ("BACKGROUND", (0,0), (-1,0), colors.lightgrey),
        ]))
        elements.append(t)

    # Border drawing
    def add_page_frame(canvas_obj, doc_obj):
        width, height = A4
        canvas_obj.saveState()
        canvas_obj.setLineWidth(1)
        canvas_obj.rect(30, 30, width - 60, height - 60)
        canvas_obj.restoreState()

    doc.build(elements, onFirstPage=add_page_frame, onLaterPages=add_page_frame)
    buffer.seek(0)

    return send_file(
        buffer,
        as_attachment=True,
        download_name="NRSC_Farmer_Advisor_Report.pdf",
        mimetype="application/pdf",
    )

# ---------- Download as Word ----------
@app.route("/download/word")
def download_word():
    farmer = session.get("farmer", {})
    crop = session.get("crop", {})
    averages = session.get("averages", {})
    calculated = session.get("calculated", {})
    irrigation_schedule = session.get("irrigation_schedule", [])
    fert_schedule = session.get("fert_schedule", [])

    doc = Document()
    doc.add_heading("NRSC Farmer Advisor - Soil & Crop Report", level=1)

    doc.add_heading("Farmer Details", level=2)
    doc.add_paragraph(f"Name: {farmer.get('name', '')}")
    doc.add_paragraph(f"Mobile: {farmer.get('mobile', '')}")
    doc.add_paragraph(f"Village: {farmer.get('village', '')}")
    doc.add_paragraph(f"Mandal: {farmer.get('mandal', '')}")
    doc.add_paragraph(f"District: {farmer.get('district', '')}")

    doc.add_heading("Crop & Soil Details", level=2)
    doc.add_paragraph(f"GPS Location: {crop.get('gps_location', '')}")
    doc.add_paragraph(f"Address: {crop.get('gps_address', '')}")
    doc.add_paragraph(f"Soil Type: {crop.get('soil_type', '')}")
    doc.add_paragraph(f"Crop: {crop.get('crop', '')}")
    doc.add_paragraph(f"Variety: {crop.get('variety', '')}")
    doc.add_paragraph(f"Season: {crop.get('season', '')}")
    doc.add_paragraph(f"Crop Duration (days): {crop.get('duration', '')}")
    doc.add_paragraph(f"Recommended N:P:K (Excel): {crop.get('recommended_npk', '')}")
    area_ha = float(crop.get("land_area", 1.0) or 1.0)
    doc.add_paragraph(f"Field / Land Area (ha): {area_ha:.2f}")
    ty_q = crop.get("target_yield_q")
    if ty_q is not None:
        doc.add_paragraph(f"Target Yield for STCR (q/ha): {ty_q}")
    doc.add_paragraph(f"Date of Sowing: {crop.get('sowing_date', '')}")

    doc.add_heading("Average Sensor Readings", level=2)
    for label, key in [
        ("Soil Nitrogen (N)", "soil_n"),
        ("Soil Phosphorus (P)", "soil_p"),
        ("Soil Potassium (K)", "soil_k"),
        ("Electrical Conductivity (EC)", "ec"),
        ("Soil Temperature (°C)", "soil_temp"),
        ("Soil pH", "ph"),
        ("Soil Moisture (%)", "soil_moisture"),
        ("Air Temperature (°C)", "air_temp"),
        ("Air Humidity (%)", "air_humidity"),
        ("Water/Soil TDS", "tds"),
    ]:
        doc.add_paragraph(f"{label}: {averages.get(key, 0):.2f}")

    doc.add_heading("Fertilizer Recommendation (STCR-based)", level=2)
    doc.add_paragraph(
        f"Final N:P:K recommendation (kg/ha): "
        f"{calculated.get('recommended_npk_after', '')}"
    )
    doc.add_paragraph(
        f"(Per ha) {calculated.get('fertilizer_form', '')}"
    )

    ud_ha = calculated.get("urea_dap_mop_per_ha", {})
    ud_field = calculated.get("urea_dap_mop_field", {})
    us_ha = calculated.get("urea_ssp_mop_per_ha", {})
    us_field = calculated.get("urea_ssp_mop_field", {})

    doc.add_paragraph("Option 1: Urea + DAP + MoP")
    doc.add_paragraph(
        f"  Per ha: Urea {ud_ha.get('urea',0):.2f} kg, "
        f"DAP {ud_ha.get('dap',0):.2f} kg, "
        f"MoP {ud_ha.get('mop',0):.2f} kg"
    )
    doc.add_paragraph(
        f"  For field ({area_ha:.2f} ha): Urea {ud_field.get('urea',0):.2f} kg, "
        f"DAP {ud_field.get('dap',0):.2f} kg, "
        f"MoP {ud_field.get('mop',0):.2f} kg"
    )

    doc.add_paragraph("Option 2: Urea + SSP + MoP")
    doc.add_paragraph(
        f"  Per ha: Urea {us_ha.get('urea',0):.2f} kg, "
        f"SSP {us_ha.get('ssp',0):.2f} kg, "
        f"MoP {us_ha.get('mop',0):.2f} kg"
    )
    doc.add_paragraph(
        f"  For field ({area_ha:.2f} ha): Urea {us_field.get('urea',0):.2f} kg, "
        f"SSP {us_field.get('ssp',0):.2f} kg, "
        f"MoP {us_field.get('mop',0):.2f} kg"
    )

    doc.add_heading("Fertilizer split schedule", level=2)
    for fs in fert_schedule:
        doc.add_paragraph(
            f"{fs['stage']} | {fs['das']} DAS | {fs['date']} | "
            f"N: {fs['N']} kg/ha, P2O5: {fs['P2O5']} kg/ha, K2O: {fs['K2O']} kg/ha"
        )
        doc.add_paragraph(
            f"    Urea + DAP + MoP (field): "
            f"Urea {fs['urea_dap_urea_field']} kg, "
            f"DAP {fs['urea_dap_dap_field']} kg, "
            f"MoP {fs['urea_dap_mop_field']} kg"
        )
        doc.add_paragraph(
            f"    Urea + SSP + MoP (field): "
            f"Urea {fs['urea_ssp_urea_field']} kg, "
            f"SSP {fs['urea_ssp_ssp_field']} kg, "
            f"MoP {fs['urea_ssp_mop_field']} kg"
        )

    doc.add_heading("Irrigation Schedule", level=2)
    if irrigation_schedule:
        doc.add_paragraph(f"Total irrigations: {len(irrigation_schedule)}")
        for irr in irrigation_schedule:
            doc.add_paragraph(
                f"Irrigation {irr['no']}: {irr['date']} | "
                f"Depth {irr['depth_mm']:.1f} mm | "
                f"Water {irr['water_m3_per_ha']:.1f} m3/ha | "
                f"{irr['water_m3_field']:.1f} m3 for field"
            )
    else:
        doc.add_paragraph("No irrigation schedule generated.")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return send_file(
        buffer,
        as_attachment=True,
        download_name="NRSC_Farmer_Advisor_Report.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

# ---------- OPTIONAL: Simple DB test route ----------
@app.route("/db-test")
def db_test():
    if engine is None:
        return "DATABASE_URL not set or DB connection not available", 500
    try:
        with engine.connect() as conn:
            value = conn.execute(text("SELECT 'neon_connected'")).scalar_one()
        return f"DB OK: {value}"
    except Exception as e:
        return f"DB error: {e}", 500


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port, debug=False)
